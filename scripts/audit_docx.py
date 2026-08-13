#!/usr/bin/env python3
"""Read-only structural audit for a DOCX package."""

import argparse
import json
import re
import sys
from pathlib import Path
from zipfile import ZipFile, BadZipFile

from docx import Document
from lxml import etree


JP_RE = re.compile(r"[\u3040-\u30ff\u3400-\u9fff]")
STEP_RE = re.compile(r"^[A-Z]\d+\.$")
CLIENT_NAME_RE = re.compile(r"\b(?:Microsoft|Google)\b", re.IGNORECASE)
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def audit(path: Path) -> dict:
    result = {"file": str(path.resolve()), "size": path.stat().st_size}
    try:
        with ZipFile(path) as package:
            names = package.namelist()
            result["zip_integrity"] = package.testzip() is None
            result["package_parts"] = len(names)
            result["media_parts"] = sum(n.startswith("word/media/") for n in names)
            result["diagram_parts"] = sum(n.startswith("word/diagrams/") for n in names)
            result["embedding_parts"] = sum(n.startswith("word/embeddings/") for n in names)
            document_root = etree.fromstring(package.read("word/document.xml"))
            numbered = document_root.xpath(".//w:tc[1]/w:p[w:pPr/w:numPr]", namespaces=NS)
            result["automatic_number_cells"] = len(numbered)
            conflicts = []
            for paragraph in numbered:
                visible = "".join(paragraph.xpath(".//w:t/text()", namespaces=NS)).strip()
                if STEP_RE.fullmatch(visible):
                    conflicts.append(visible)
            result["automatic_plus_typed_conflicts"] = conflicts
            result["numbering_part_present"] = "word/numbering.xml" in names
            client_candidates = {}
            for name in names:
                if not (name.startswith("word/") and name.endswith(".xml")):
                    continue
                try:
                    part_root = etree.fromstring(package.read(name))
                except etree.XMLSyntaxError:
                    continue
                visible_text = " ".join(
                    node.text or ""
                    for node in part_root.iter()
                    if etree.QName(node).localname in {"t", "instrText", "delText"}
                )
                hits = CLIENT_NAME_RE.findall(visible_text)
                if hits:
                    client_candidates[name] = sorted(set(hits), key=str.casefold)
            result["client_name_candidates"] = client_candidates
    except BadZipFile:
        result["zip_integrity"] = False
        return result

    doc = Document(path)
    paragraphs = list(doc.paragraphs)
    table_paragraphs = [p for t in doc.tables for r in t.rows for c in r.cells for p in c.paragraphs]
    all_paragraphs = paragraphs + table_paragraphs
    result.update(
        tables=len(doc.tables),
        inline_shapes=len(doc.inline_shapes),
        body_paragraphs=len(paragraphs),
        table_paragraphs=len(table_paragraphs),
        japanese_paragraphs=sum(bool(JP_RE.search(p.text)) for p in all_paragraphs),
    )

    step_ids = []
    for table in doc.tables:
        for row in table.rows:
            if row.cells:
                value = row.cells[0].text.strip()
                if STEP_RE.fullmatch(value):
                    step_ids.append(value)
    result["explicit_step_ids"] = len(step_ids)
    result["first_step_id"] = step_ids[0] if step_ids else None
    result["last_step_id"] = step_ids[-1] if step_ids else None
    result["numbering_mode"] = (
        "mixed-conflict" if result.get("automatic_plus_typed_conflicts")
        else "automatic" if result.get("automatic_number_cells")
        else "explicit" if step_ids
        else "none-detected"
    )
    return result


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", nargs="+", type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    reports = [audit(p) for p in args.docx]
    if args.json:
        print(json.dumps(reports, ensure_ascii=False, indent=2))
    else:
        for report in reports:
            print("\n" + report["file"])
            for key, value in report.items():
                if key != "file":
                    print(f"  {key}: {value}")


if __name__ == "__main__":
    main()
